"""
Creating and Styling Word Documents
This demonstrates creating .docx files with formatting
"""

import docx

# Creating a new document
doc = docx.Document()

# Adding paragraphs
doc.add_paragraph('This is a normal paragraph.')

# Adding styled paragraphs
doc.add_paragraph('This is a quote.', 'Intense Quote')

# Adding headings (levels 0-4)
doc.add_heading('Heading Level 1', 0)
doc.add_heading('Heading Level 2', 1)
doc.add_heading('Heading Level 3', 2)

# Adding paragraph with runs (for styling)
para = doc.add_paragraph('Normal text, ')
run = para.add_run('bold text, ')
run.bold = True
run = para.add_run('italic text, ')
run.italic = True
run = para.add_run('underlined text.')
run.underline = True

# Adding page breaks
doc.add_page_break()

# Adding pictures
doc.add_picture('image.png', width=docx.shared.Inches(2))

# Saving document
doc.save('created_document.docx')
print('Document created successfully')
