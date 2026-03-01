"""
Reading Word Documents with python-docx
This demonstrates reading .docx files and extracting text
"""

import docx

# Opening a Word document
doc = docx.Document('example.docx')

# Getting all paragraphs
print(f'Number of paragraphs: {len(doc.paragraphs)}')

# Reading paragraph text
for para in doc.paragraphs:
    print(para.text)

# Accessing runs within a paragraph
# A run is continuous text with same style
para = doc.paragraphs[0]
print(f'\nFirst paragraph has {len(para.runs)} runs')

for run in para.runs:
    print(f'Run text: {run.text}')
    print(f'  Bold: {run.bold}')
    print(f'  Italic: {run.italic}')
    print(f'  Underline: {run.underline}')

# Getting full text of document
fullText = []
for para in doc.paragraphs:
    fullText.append(para.text)
text = '\n'.join(fullText)
print('\n' + text)
